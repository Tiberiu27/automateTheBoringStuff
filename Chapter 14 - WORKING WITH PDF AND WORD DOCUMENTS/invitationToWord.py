#!python3
#invitationToWord.py - Custom Invitations as Words Documents
import docx, sys, os

abspath = os.path.abspath(sys.argv[0])
dname = os.path.dirname(abspath)
os.chdir(dname)

doc = docx.Document('guests.docx')
guestList = open('guests.txt')

for guest in guestList.readlines():
    print(guest)
    paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of')
    paraObj1.style = 'custom1'
    paraObj2 = doc.add_paragraph(guest)
    paraObj2.bold = True
    paraObj3 = doc.add_paragraph('at 1010 Memory Lane on the evening of')
    paraObj3.style = 'custom1'
    doc.add_paragraph('April 1st')
    paraObj4 = doc.add_paragraph("at 7 o'clock")
    paraObj4.style = 'custom1'
    paraObj4.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

for paragraph in doc.paragraphs:
    paragraph.alignment = 1

doc.save('guests.docx')
